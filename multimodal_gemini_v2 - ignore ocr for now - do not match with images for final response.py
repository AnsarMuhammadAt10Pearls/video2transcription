import os
import cv2
import base64
import tempfile
import hashlib
import time
import json
import pytesseract
import sys
from datetime import datetime, timedelta
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image
import io
from pathlib import Path
import numpy as np
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
import docx
from docx.shared import Inches
import argparse
import glob

# Configure stdout to use UTF-8 if possible (helps with Windows encoding issues)
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        # For older Python versions
        pass

# Configure Tesseract OCR path
pytesseract.pytesseract.tesseract_cmd = r'D:\tesseract\tesseract.exe'

# Load environment variables
load_dotenv()

# Constants for the application
DEFAULT_QUERY = "Complete Journey of Committing and Deploying of a Symbox Component"
ADDITIONAL_PROMPTS = [
    "How do you integrate an external system with Symbox?"
]
VIDEOS_FOLDER = "videos"
FRAME_EXTRACTION_INTERVAL = 10  # Extract a frame every 10 seconds
OCR_CACHE_DIR = "ocr_cache"
TEMP_DIR = "temp_frames"
VECTOR_DB_PATH = "vectordb"
INDEX_CACHE_DIR = "index_cache"

# Model settings
MODEL_NAME = "gemini-1.5-flash"  # Using Gemini-1.5-flash, can switch to gemini-2.0 when available
TEMPERATURE = 0.3
MAX_TOKENS = 4000

# Configure Gemini client
gemini_api_key = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=gemini_api_key)

# Initialize safety settings to allow dangerous deserialization
safety_settings = {
    genai.types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: genai.types.HarmBlockThreshold.BLOCK_NONE,
    genai.types.HarmCategory.HARM_CATEGORY_HATE_SPEECH: genai.types.HarmBlockThreshold.BLOCK_NONE,
    genai.types.HarmCategory.HARM_CATEGORY_HARASSMENT: genai.types.HarmBlockThreshold.BLOCK_NONE,
    genai.types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: genai.types.HarmBlockThreshold.BLOCK_NONE,
}

# Initialize embeddings model
embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Create directories if they don't exist
os.makedirs(OCR_CACHE_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(INDEX_CACHE_DIR, exist_ok=True)

# Store token usage statistics
token_usage = {
    "total_input_tokens": 0,
    "total_output_tokens": 0,
    "requests": 0
}

# Track performance metrics
performance_metrics = {
    "video_processing_times": {},
    "total_runtime": 0,
    "video_lengths": {}
}

def get_video_duration(video_path):
    """
    Get the duration of a video in seconds
    
    Args:
        video_path: Path to the video file
    
    Returns:
        Duration in seconds
    """
    video = cv2.VideoCapture(video_path)
    fps = video.get(cv2.CAP_PROP_FPS)
    frame_count = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
    video.release()
    
    duration = frame_count / fps if fps > 0 else 0
    return duration

def get_video_hash(video_path):
    """
    Generate a hash for a video file to use as a cache key
    
    Args:
        video_path: Path to the video file
    
    Returns:
        Hash string for the video
    """
    # Use file size and modification time as a proxy for content hash
    # This is faster than hashing the entire file content
    file_stats = os.stat(video_path)
    hash_input = f"{video_path}_{file_stats.st_size}_{file_stats.st_mtime}"
    return hashlib.md5(hash_input.encode()).hexdigest()

def perform_ocr(image):
    """
    Perform OCR on an image using Tesseract
    
    Args:
        image: PIL Image object
    
    Returns:
        Extracted text from the image
    """
    # Convert PIL image to OpenCV format if needed
    if isinstance(image, Image.Image):
        image = np.array(image)
        image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)
    
    # Convert to grayscale for better OCR results
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Apply some preprocessing to improve OCR results
    gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
    
    # Perform OCR
    text = pytesseract.image_to_string(gray)
    return text.strip()

def cache_ocr_results(image, video_name, frame_idx):
    """
    Perform OCR and cache the results
    
    Args:
        image: PIL Image object
        video_name: Name of the video file
        frame_idx: Frame index
    
    Returns:
        OCR text
    """
    # Create a cache key
    video_hash = get_video_hash(video_name)
    cache_key = f"{video_hash}_{frame_idx}"
    cache_path = os.path.join(OCR_CACHE_DIR, f"{cache_key}.txt")
    
    # Check if result is already cached
    if os.path.exists(cache_path):
        with open(cache_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    # Perform OCR
    ocr_text = perform_ocr(image)
    
    # Cache the result
    with open(cache_path, 'w', encoding='utf-8') as f:
        f.write(ocr_text)
    
    return ocr_text

def extract_frames_from_video(video_path, interval=FRAME_EXTRACTION_INTERVAL, force_reprocess=False):
    """
    Extract frames from a video at specified intervals
    
    Args:
        video_path: Path to the video file
        interval: Interval in seconds between frames to extract
        force_reprocess: Force reprocessing even if cached
    
    Returns:
        List of extracted frames as PIL Images, frame timestamps, and OCR text
    """
    frames = []
    frame_timestamps = []
    ocr_results = []
    
    # Track processing time
    start_time = time.time()
    
    # Check if we have a cached version
    video_hash = get_video_hash(video_path)
    cache_file = os.path.join(INDEX_CACHE_DIR, f"{video_hash}_frames.json")
    
    if os.path.exists(cache_file) and not force_reprocess:
        print(f"Using cached frames for {video_path}")
        with open(cache_file, 'r', encoding='utf-8') as f:
            cache_data = json.load(f)
            
        # Load frames from cached paths
        for frame_info in cache_data['frames']:
            frame_path = frame_info['path']
            if os.path.exists(frame_path):
                frames.append(Image.open(frame_path))
                frame_timestamps.append(frame_info['timestamp'])
                ocr_results.append(frame_info['ocr_text'])
        
        # Get video duration
        video_duration = get_video_duration(video_path)
        performance_metrics["video_lengths"][video_path] = video_duration
            
        if frames:
            # Record processing time
            end_time = time.time()
            processing_time = end_time - start_time
            performance_metrics["video_processing_times"][video_path] = processing_time
            
            return frames, frame_timestamps, ocr_results
    
    print(f"Extracting frames from {video_path}")
    
    # Open the video file
    video = cv2.VideoCapture(video_path)
    
    # Get video properties
    fps = video.get(cv2.CAP_PROP_FPS)
    frame_count = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
    duration = frame_count / fps
    
    # Store video duration
    performance_metrics["video_lengths"][video_path] = duration
    
    # Calculate frame indices to extract based on interval
    frame_indices = [int(fps * i) for i in range(0, int(duration), interval)]
    
    # Prepare cache data
    cache_data = {
        'video_path': video_path,
        'hash': video_hash,
        'frames': []
    }
    
    for i, frame_idx in enumerate(frame_indices):
        video.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
        success, frame = video.read()
        
        if success:
            # Convert BGR to RGB (OpenCV uses BGR, PIL uses RGB)
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            pil_image = Image.fromarray(rgb_frame)
            
            # Save frame to temp directory
            base_name = Path(video_path).stem
            frame_path = os.path.join(TEMP_DIR, f"{base_name}_frame_{i:04d}.jpg")
            pil_image.save(frame_path)
            
            # Perform OCR
            ocr_text = cache_ocr_results(pil_image, video_path, i)
            
            # Add frame to our collection with its timestamp
            frames.append(pil_image)
            timestamp = frame_idx / fps
            frame_timestamps.append(timestamp)
            ocr_results.append(ocr_text)
            
            # Add to cache data
            cache_data['frames'].append({
                'path': frame_path,
                'timestamp': timestamp,
                'ocr_text': ocr_text
            })
    
    video.release()
    
    # Save cache data
    with open(cache_file, 'w', encoding='utf-8') as f:
        json.dump(cache_data, f, ensure_ascii=False)
    
    # Record processing time
    end_time = time.time()
    processing_time = end_time - start_time
    performance_metrics["video_processing_times"][video_path] = processing_time
    
    print(f"Extracted {len(frames)} frames from {video_path}")
    return frames, frame_timestamps, ocr_results

def save_frames_to_temp(frames, video_name):
    """
    Save frames to temporary directory and return paths
    
    Args:
        frames: List of PIL Image objects
        video_name: Name of the video file
    
    Returns:
        List of paths to saved frames
    """
    paths = []
    base_name = Path(video_name).stem
    
    for i, frame in enumerate(frames):
        frame_path = os.path.join(TEMP_DIR, f"{base_name}_frame_{i:04d}.jpg")
        frame.save(frame_path)
        paths.append(frame_path)
    
    return paths

def analyze_frames_with_gemini(frames, prompt):
    """
    Analyze a batch of frames with Gemini
    
    Args:
        frames: List of PIL Image objects
        prompt: The text prompt for Gemini
    
    Returns:
        Gemini's response
    """
    # Create a Gemini model instance
    model = genai.GenerativeModel(
        MODEL_NAME,
        safety_settings=safety_settings
    )
    
    # Prepare content parts
    content_parts = [prompt]
    
    # Add frames to the content parts
    for frame in frames:
        content_parts.append(frame)
    
    # Send request to Gemini
    response = model.generate_content(
        content_parts,
        generation_config=genai.types.GenerationConfig(
            temperature=TEMPERATURE,
            max_output_tokens=MAX_TOKENS,
        ),
        safety_settings=safety_settings
    )
    
    # Update token usage statistics
    if hasattr(response, 'usage_metadata'):
        token_usage["total_input_tokens"] += response.usage_metadata.prompt_token_count
        token_usage["total_output_tokens"] += response.usage_metadata.candidates_token_count
    token_usage["requests"] += 1
    
    return response.text

def create_vector_store_from_frames(frame_analysis_results, video_names, frame_timestamps, ocr_results):
    """
    Create a vector store from frame analysis results
    
    Args:
        frame_analysis_results: List of Claude's analysis of frames
        video_names: Names of the source videos
        frame_timestamps: Timestamps for frames in seconds
        ocr_results: OCR text extracted from frames
    
    Returns:
        FAISS vector store
    """
    documents = []
    flat_ocr_results = []
    
    # Flatten the OCR results from multiple videos
    for video_ocr in ocr_results:
        flat_ocr_results.extend(video_ocr)
    
    # Create documents from frame analysis results
    for i, analysis in enumerate(frame_analysis_results):
        # Determine which video and timestamp this came from
        video_index = i // len(frame_timestamps[0])
        frame_index = i % len(frame_timestamps[0])
        
        if video_index < len(video_names) and frame_index < len(frame_timestamps[video_index]):
            video_name = video_names[video_index]
            timestamp = frame_timestamps[video_index][frame_index]
            
            # Get corresponding OCR text
            ocr_text = ""
            if i < len(flat_ocr_results):
                ocr_text = flat_ocr_results[i]
            
            # Format timestamp as minutes:seconds
            timestamp_formatted = f"{int(timestamp // 60)}:{int(timestamp % 60):02d}"
            
            # Combine AI analysis with OCR text
            combined_content = analysis
            if ocr_text:
                combined_content = f"{analysis}\n\nOCR Text:\n{ocr_text}"
            
            # Create document with metadata
            doc = Document(
                page_content=combined_content,
                metadata={
                    "source": video_name,
                    "timestamp": timestamp,
                    "timestamp_formatted": timestamp_formatted,
                    "ocr_text": ocr_text
                }
            )
            
            documents.append(doc)
    
    # Create vector store
    vector_store = FAISS.from_documents(documents, embeddings_model)
    
    # Save the vector store
    vector_store.save_local(VECTOR_DB_PATH)
    
    return vector_store

def generate_final_response_with_gemini(query, vector_store, video_frames, all_video_names, all_frame_timestamps):
    """
    Generate a final response to the query using RAG with Gemini
    
    Args:
        query: The user query
        vector_store: FAISS vector store of frame analyses
        video_frames: List of video frames as PIL Images
        all_video_names: List of video file names
        all_frame_timestamps: List of timestamps for frames
    
    Returns:
        Gemini's response, key frames used, and relevant frames
    """
    # Retrieve relevant documents using combined search
    docs = combined_search(query, vector_store)
    
    # Prepare content for Gemini
    context_items = []
    for doc in docs:
        source_video = os.path.basename(doc.metadata['source'])
        timestamp = doc.metadata['timestamp_formatted']
        context_items.append(
            f"Source: {source_video} at {timestamp}\n{doc.page_content}"
        )
    
    context = "\n\n".join(context_items)
    
    # Track relevant frames from search results
    relevant_frames = []
    relevant_frame_info = []
    
    for doc in docs:
        video_name = doc.metadata['source']
        timestamp = doc.metadata['timestamp']
        
        # Find which video this is from
        video_index = -1
        for i, name in enumerate(all_video_names):
            if name == video_name:
                video_index = i
                break
        
        if video_index >= 0:
            # Find the closest frame to this timestamp
            timestamps = all_frame_timestamps[video_index]
            closest_frame_idx = min(range(len(timestamps)), 
                              key=lambda i: abs(timestamps[i] - timestamp))
            
            # Calculate the absolute index in the combined frames list
            absolute_idx = sum(len(ts) for ts in all_frame_timestamps[:video_index]) + closest_frame_idx
            
            if 0 <= absolute_idx < len(video_frames):
                relevant_frames.append(video_frames[absolute_idx])
                relevant_frame_info.append({
                    'video': video_name,
                    'timestamp': timestamp,
                    'timestamp_formatted': doc.metadata['timestamp_formatted']
                })
    
    # Select a few key frames to include with the prompt
    key_frame_indices = [0, min(2, len(relevant_frames)-1)]
    key_frames = [relevant_frames[i] for i in key_frame_indices if i < len(relevant_frames)]
    
    # Create prompt for Gemini
    prompt = f"""
You are a knowledgeable assistant specializing in Symbox platform development workflows.

I need a comprehensive explanation about: {query}

Based on analysis of video tutorials, here's relevant information:

{context}

Please provide a detailed explanation that answers the user's question.
Structure your answer with clear headings and steps where appropriate.
IMPORTANT: Always include video references and timestamps for each piece of information in your answer.
For each point or step you explain, add a citation in the format "[Source: Video_Name at MM:SS]" at the end of the relevant paragraph or bullet point.
"""

    # Create a Gemini model instance
    model = genai.GenerativeModel(
        MODEL_NAME,
        safety_settings=safety_settings
    )
    
    # Prepare content parts
    content_parts = [prompt]
    
    # Add key frames to the content parts
    for frame in key_frames:
        content_parts.append(frame)
    
    # Send request to Gemini
    response = model.generate_content(
        content_parts,
        generation_config=genai.types.GenerationConfig(
            temperature=TEMPERATURE,
            max_output_tokens=MAX_TOKENS,
        ),
        safety_settings=safety_settings
    )
    
    # Update token usage statistics
    if hasattr(response, 'usage_metadata'):
        token_usage["total_input_tokens"] += response.usage_metadata.prompt_token_count
        token_usage["total_output_tokens"] += response.usage_metadata.candidates_token_count
    token_usage["requests"] += 1
    
    return response.text, key_frames, relevant_frames, relevant_frame_info

def save_response_as_docx(response_text, key_frames, relevant_frames, relevant_frame_info, query):
    """
    Save the response text and frames as a docx file
    
    Args:
        response_text: Text response from Claude
        key_frames: List of key frames as PIL Images
        relevant_frames: List of relevant frames from search results
        relevant_frame_info: Metadata about relevant frames
        query: The original query
    
    Returns:
        Path to the saved docx file
    """
    # Create a new document
    doc = docx.Document()
    
    # Set document properties
    doc.core_properties.title = f"Symbox Guide: {query}"
    doc.core_properties.author = "Symbox Documentation Generator"
    doc.core_properties.created = datetime.now()
    
    # Add a cover page
    doc.add_heading(f"Symbox Documentation", level=0)
    doc.add_heading(query, level=1)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Add a page break after the cover
    doc.add_page_break()
    
    # Add table of contents heading
    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    run = toc_para.add_run("(Right-click and select 'Update Field' to update the table of contents)")
    run.italic = True
    run.font.size = docx.shared.Pt(9)
    
    # Add TOC field
    para = doc.add_paragraph()
    run = para.add_run()
    fld_char = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar', 
                                        {'fldCharType': 'begin'})
    run._element.append(fld_char)
    
    instrText = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText', {})
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)
    
    fld_char = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar', 
                                        {'fldCharType': 'separate'})
    run._element.append(fld_char)
    
    # Add static TOC placeholder text
    run = para.add_run("Table of Contents placeholder. Right-click and select 'Update Field' to update.")
    
    fld_char = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar', 
                                        {'fldCharType': 'end'})
    run._element.append(fld_char)
    
    # Add a page break after TOC
    doc.add_page_break()
    
    # Add reference section before content
    doc.add_heading("Video References", level=1)
    
    # Create a set of unique video files used in this response
    referenced_videos = set()
    for info in relevant_frame_info:
        referenced_videos.add(os.path.basename(info['video']))
    
    # Add list of video references
    for video in sorted(referenced_videos):
        doc.add_paragraph(f"• {video}", style='List Bullet')
    
    # Add a page break after references
    doc.add_page_break()
    
    # Add content heading
    doc.add_heading("Guide Content", level=1)
    
    # Process the response text to handle markdown formatting
    paragraphs = response_text.split('\n\n')
    i = 0
    
    while i < len(paragraphs):
        para = paragraphs[i].strip()
        
        # Skip empty paragraphs
        if not para:
            i += 1
            continue
        
        # Check if paragraph is a heading (starts with #)
        if para.startswith('#'):
            # Count the number of # to determine heading level
            heading_level = 0
            for char in para:
                if char == '#':
                    heading_level += 1
                else:
                    break
            
            heading_level = min(heading_level, 9)  # Word only supports heading levels 1-9
            heading_text = para[heading_level:].strip()
            doc.add_heading(heading_text, level=heading_level)
        
        # Check if paragraph is a bulleted list (starts with * or -)
        elif para.startswith('* ') or para.startswith('- '):
            # Handle multi-line bulleted list
            bullet_items = [para]
            j = i + 1
            while j < len(paragraphs) and (paragraphs[j].strip().startswith('* ') or 
                                          paragraphs[j].strip().startswith('- ')):
                bullet_items.append(paragraphs[j].strip())
                j += 1
            
            for item in bullet_items:
                # Remove the bullet marker and add as a list item
                item_text = item[2:].strip()
                p = doc.add_paragraph(item_text, style='List Bullet')
                
                # Make citation in italics if present
                if "[Source:" in item_text:
                    runs = p.runs
                    if runs:
                        text = runs[-1].text
                        source_idx = text.find("[Source:")
                        if source_idx >= 0:
                            # Split the run at the citation
                            runs[-1].text = text[:source_idx]
                            citation_run = p.add_run(text[source_idx:])
                            citation_run.italic = True
                            citation_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
            
            i = j - 1  # Update the counter to skip processed items
        
        # Check if paragraph is a numbered list (starts with 1., 2., etc.)
        elif len(para) > 1 and para[0].isdigit() and para[1:].startswith('. '):
            # Handle multi-line numbered list
            number_items = [para]
            current_number = int(para[0])
            j = i + 1
            
            while j < len(paragraphs):
                next_para = paragraphs[j].strip()
                if next_para and next_para.startswith(f"{current_number + 1}. "):
                    number_items.append(next_para)
                    current_number += 1
                    j += 1
                else:
                    break
            
            for item in number_items:
                # Find the position after the number and period
                pos = item.find('. ')
                if pos != -1:
                    item_text = item[pos + 2:].strip()
                    p = doc.add_paragraph(item_text, style='List Number')
                    
                    # Make citation in italics if present
                    if "[Source:" in item_text:
                        runs = p.runs
                        if runs:
                            text = runs[-1].text
                            source_idx = text.find("[Source:")
                            if source_idx >= 0:
                                # Split the run at the citation
                                runs[-1].text = text[:source_idx]
                                citation_run = p.add_run(text[source_idx:])
                                citation_run.italic = True
                                citation_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
            
            i = j - 1  # Update the counter to skip processed items
            
        # Regular paragraph
        else:
            p = doc.add_paragraph(para)
            
            # Make citation in italics if present
            if "[Source:" in para:
                runs = p.runs
                if runs:
                    text = runs[-1].text
                    source_idx = text.find("[Source:")
                    if source_idx >= 0:
                        # Split the run at the citation
                        runs[-1].text = text[:source_idx]
                        citation_run = p.add_run(text[source_idx:])
                        citation_run.italic = True
                        citation_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
        
        i += 1
    
    # Add a page break before images
    doc.add_page_break()
    
    # Add a section for relevant frames from the search with timestamps
    if relevant_frames:
        doc.add_heading("Relevant Tutorial Screenshots", level=1)
        
        for i, (frame, info) in enumerate(zip(relevant_frames, relevant_frame_info)):
            # Save the image to a temporary file
            temp_img_path = os.path.join(TEMP_DIR, f"relevant_frame_{i}.jpg")
            frame.save(temp_img_path)
            
            # Add image to document with caption
            video_name = os.path.basename(info['video'])
            caption = f"Source: {video_name} at {info['timestamp_formatted']}"
            
            # Add caption in a distinct format
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(caption)
            caption_run.bold = True
            caption_run.font.size = docx.shared.Pt(10)
            
            # Add the image
            doc.add_picture(temp_img_path, width=Inches(6))
            
            # Add a small space between images
            doc.add_paragraph()
    
    # Add a section for key frames used in the prompt
    doc.add_heading("Additional Tutorial Screenshots", level=1)
    
    # Save key frames to temporary files and add to document
    for i, frame in enumerate(key_frames):
        # Save the image to a temporary file
        temp_img_path = os.path.join(TEMP_DIR, f"key_frame_{i}.jpg")
        frame.save(temp_img_path)
        
        # Add image to document with caption
        doc.add_paragraph(f"Frame {i+1}")
        doc.add_picture(temp_img_path, width=Inches(6))
    
    # Save the document
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    docx_path = f"response_{timestamp}.docx"
    doc.save(docx_path)
    
    return docx_path

def search_by_ocr(query, vector_store):
    """
    Search for frames that contain OCR text matching the query
    
    Args:
        query: The search query
        vector_store: FAISS vector store
    
    Returns:
        List of matching documents
    """
    # Extract the underlying docstore from the vector store
    docstore = vector_store.docstore
    
    # Search for matches in OCR text
    matches = []
    for doc_id, doc in docstore._dict.items():
        if 'ocr_text' in doc.metadata and doc.metadata['ocr_text']:
            # Simple substring match in OCR text
            if query.lower() in doc.metadata['ocr_text'].lower():
                matches.append(doc)
    
    return matches

def combined_search(query, vector_store):
    """
    Perform both semantic search and OCR text search
    
    Args:
        query: The search query
        vector_store: FAISS vector store
    
    Returns:
        Combined list of relevant documents
    """
    # Semantic search
    semantic_results = vector_store.similarity_search(query, k=3)
    
    # OCR text search
    ocr_results = search_by_ocr(query, vector_store)
    
    # Combine results, removing duplicates
    combined_results = semantic_results.copy()
    seen_ids = {doc.metadata['source'] + str(doc.metadata['timestamp']) for doc in combined_results}
    
    for doc in ocr_results:
        doc_id = doc.metadata['source'] + str(doc.metadata['timestamp'])
        if doc_id not in seen_ids:
            combined_results.append(doc)
            seen_ids.add(doc_id)
    
    return combined_results

def main():
    """Main function to orchestrate the multimodal RAG workflow"""
    # Track total program runtime
    program_start_time = time.time()
    
    parser = argparse.ArgumentParser(description='Multimodal RAG analysis on video files')
    parser.add_argument('--force-reprocess', action='store_true', help='Force reprocessing of videos even if cached')
    args = parser.parse_args()
    
    # Find all video files in the videos folder
    video_files = []
    for ext in ['*.mp4', '*.avi', '*.mov', '*.mkv']:
        video_files.extend(glob.glob(os.path.join(VIDEOS_FOLDER, ext)))
    
    if not video_files:
        print(f"No video files found in {VIDEOS_FOLDER} directory!")
        return
    
    # Print configuration information
    print("=" * 50)
    print("MODEL CONFIGURATION:")
    print(f"LLM Model: {MODEL_NAME} (Google Gemini)")
    print(f"Temperature: {TEMPERATURE}")
    print(f"Max output tokens: {MAX_TOKENS}")
    print(f"Embedding Algorithm: {embeddings_model.model_name}")
    print(f"Vector DB: FAISS")
    print(f"Index Path: {VECTOR_DB_PATH}")
    print(f"OCR Engine: Tesseract (Path: {pytesseract.pytesseract.tesseract_cmd})")
    print("=" * 50)
    
    print(f"Starting multimodal RAG analysis on {len(video_files)} video files...")
    
    all_frames = []
    all_video_names = []
    all_frame_timestamps = []
    all_ocr_results = []
    
    # Check if we can load a previously saved vector store
    vector_store = None
    if os.path.exists(VECTOR_DB_PATH) and not args.force_reprocess:
        try:
            print("Loading existing vector store...")
            vector_store = FAISS.load_local(VECTOR_DB_PATH, embeddings_model)
            
            # We still need to load frames for visualization
            for video_file in video_files:
                print(f"Loading frames from: {video_file}")
                frames, timestamps, ocr_texts = extract_frames_from_video(video_file, force_reprocess=False)
                all_frames.extend(frames)
                all_video_names.append(video_file)
                all_frame_timestamps.append(timestamps)
                all_ocr_results.append(ocr_texts)
        except Exception as e:
            print(f"Error loading vector store: {e}")
            vector_store = None
    
    # If we couldn't load the vector store, process videos and create it
    if vector_store is None:
        # Extract frames from all videos
        for video_file in video_files:
            print(f"Processing video: {video_file}")
            frames, timestamps, ocr_texts = extract_frames_from_video(video_file, force_reprocess=args.force_reprocess)
            all_frames.extend(frames)
            all_video_names.append(video_file)
            all_frame_timestamps.append(timestamps)
            all_ocr_results.append(ocr_texts)
        
        # Analyze frames in batches (Gemini can handle multiple images per request)
        frame_analysis_results = []
        batch_size = 5
        
        for i in range(0, len(all_frames), batch_size):
            print(f"Analyzing batch {i // batch_size + 1}/{(len(all_frames) + batch_size - 1) // batch_size}")
            batch_frames = all_frames[i:i+batch_size]
            
            analysis_prompt = """
            Analyze these video frames from a tutorial about Symbox component development and deployment.
            Describe what you see in detail, focusing on:
            1. Any visible user interface elements and what they represent
            2. Any text or code visible in the frame
            3. Any workflow steps being demonstrated
            4. Any version control, commit, or deployment processes shown
            
            Be specific and detailed in your observations.
            """
            
            analysis = analyze_frames_with_gemini(batch_frames, analysis_prompt)
            frame_analysis_results.append(analysis)
            
            # Save to prevent losing progress
            with open(f"frame_analysis_batch_{i // batch_size}.txt", "w", encoding='utf-8') as f:
                f.write(analysis)
            
            # Respect API rate limits
            time.sleep(1)
        
        # Create vector store from frame analyses
        print("Creating vector store from frame analyses...")
        vector_store = create_vector_store_from_frames(
            frame_analysis_results, 
            all_video_names, 
            all_frame_timestamps,
            all_ocr_results
        )
    
    # Print video processing statistics
    print("\nVIDEO PROCESSING STATISTICS:")
    for video_path, duration in performance_metrics["video_lengths"].items():
        processing_time = performance_metrics["video_processing_times"].get(video_path, 0)
        minutes = int(duration // 60)
        seconds = int(duration % 60)
        print(f"Video: {os.path.basename(video_path)}")
        print(f"  Length: {minutes} minutes {seconds} seconds")
        print(f"  Processing Time: {processing_time:.2f} seconds")
    
    # Process the default prompts automatically
    print("\nProcessing default prompts:")
    for prompt in [DEFAULT_QUERY] + ADDITIONAL_PROMPTS:
        print(f"\nGenerating response to: {prompt}")
        response, key_frames, relevant_frames, relevant_frame_info = generate_final_response_with_gemini(
            prompt, vector_store, all_frames, all_video_names, all_frame_timestamps
        )
        
        # Save the result as text file
        output_file = f"response_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        with open(output_file, "w", encoding='utf-8') as f:
            f.write(response)
        
        # Save the result as docx file with images
        docx_path = save_response_as_docx(response, key_frames, relevant_frames, relevant_frame_info, prompt)
        
        print(f"Response saved to {output_file}")
        print(f"Response with images saved to {docx_path}")
    
    # Interactive query loop
    print("\nVideo processing complete. Ready for questions!")
    print("Enter 'q' to quit.")
    
    while True:
        user_query = input("\nEnter your question about Symbox (type 'q' to quit): ")
        
        if user_query.lower() == 'q':
            print("Exiting...")
            break
        
        if not user_query.strip():
            print("Please enter a valid question.")
            continue
        
        print(f"Generating response to: {user_query}")
        response, key_frames, relevant_frames, relevant_frame_info = generate_final_response_with_gemini(
            user_query, vector_store, all_frames, all_video_names, all_frame_timestamps
        )
        
        # Save the result as text file
        output_file = f"response_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        with open(output_file, "w", encoding='utf-8') as f:
            f.write(response)
        
        # Save the result as docx file with images
        docx_path = save_response_as_docx(response, key_frames, relevant_frames, relevant_frame_info, user_query)
        
        print(f"Response saved to {output_file}")
        print(f"Response with images saved to {docx_path}")
        print("\nResponse:")
        print("-" * 80)
        print(response)
        print("-" * 80)
        
        # Calculate total program runtime
        program_end_time = time.time()
        performance_metrics["total_runtime"] = program_end_time - program_start_time
        total_runtime = timedelta(seconds=int(performance_metrics["total_runtime"]))
        
        # Print token usage for this session
        print("\nPERFORMANCE STATISTICS:")
        print(f"Total Program Runtime: {total_runtime}")
        print("\nTOKEN USAGE STATISTICS:")
        print(f"Total API Requests: {token_usage['requests']}")
        print(f"Total Input Tokens: {token_usage['total_input_tokens']}")
        print(f"Total Output Tokens: {token_usage['total_output_tokens']}")
        print(f"Total Tokens: {token_usage['total_input_tokens'] + token_usage['total_output_tokens']}")
        print("-" * 80)

if __name__ == "__main__":
    main()
