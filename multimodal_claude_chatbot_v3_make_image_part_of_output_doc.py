import os
import cv2
import base64
import tempfile
import hashlib
import time
import json
from datetime import datetime
from dotenv import load_dotenv
import anthropic
from PIL import Image
import io
from pathlib import Path
import numpy as np
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
import docx
from docx.shared import Inches

# Load environment variables
load_dotenv()

# Constants for the application
QUERY = "Complete Journey of Committing and Deploying of a Symbox Component"
VIDEO_FILES = [
    os.path.join("videos", "Product_session_17_Version_Repo_by_Eduard_02Dec2022_000.mp4"),
    os.path.join("videos", "Product_session_17_Version_Repo_by_Eduard_02Dec2022_001.mp4")
]
FRAME_EXTRACTION_INTERVAL = 10  # Extract a frame every 10 seconds
OCR_CACHE_DIR = "ocr_cache"
TEMP_DIR = "temp_frames"
VECTOR_DB_PATH = "vectordb"

# Configure Claude client
claude_api_key = os.getenv("CLAUDE_API_KEY")
claude = anthropic.Anthropic(api_key=claude_api_key)

# Initialize embeddings model
embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Create directories if they don't exist
os.makedirs(OCR_CACHE_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)

def extract_frames_from_video(video_path, interval=FRAME_EXTRACTION_INTERVAL):
    """
    Extract frames from a video at specified intervals
    
    Args:
        video_path: Path to the video file
        interval: Interval in seconds between frames to extract
    
    Returns:
        List of extracted frames as PIL Images
    """
    frames = []
    frame_timestamps = []
    
    # Open the video file
    video = cv2.VideoCapture(video_path)
    
    # Get video properties
    fps = video.get(cv2.CAP_PROP_FPS)
    frame_count = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
    duration = frame_count / fps
    
    # Calculate frame indices to extract based on interval
    frame_indices = [int(fps * i) for i in range(0, int(duration), interval)]
    
    for frame_idx in frame_indices:
        video.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
        success, frame = video.read()
        
        if success:
            # Convert BGR to RGB (OpenCV uses BGR, PIL uses RGB)
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            pil_image = Image.fromarray(rgb_frame)
            
            # Add frame to our collection with its timestamp
            frames.append(pil_image)
            timestamp = frame_idx / fps
            frame_timestamps.append(timestamp)
    
    video.release()
    
    print(f"Extracted {len(frames)} frames from {video_path}")
    return frames, frame_timestamps

def encode_image_for_claude(image):
    """
    Encode a PIL image for Claude's multimodal API
    
    Args:
        image: PIL Image object
    
    Returns:
        Base64 encoded image string
    """
    # Resize image if it's too large (Claude has limits)
    max_dim = 1500
    width, height = image.size
    
    if width > max_dim or height > max_dim:
        if width > height:
            new_width = max_dim
            new_height = int(height * (max_dim / width))
        else:
            new_height = max_dim
            new_width = int(width * (max_dim / height))
        
        image = image.resize((new_width, new_height), Image.LANCZOS)
    
    # Convert to JPEG and encode as base64
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG")
    return base64.b64encode(buffer.getvalue()).decode("utf-8")

def save_frames_to_temp(frames, video_name):
    """
    Save frames to temporary directory and return paths
    
    Args:
        frames: List of PIL Image objects
        video_name: Name of the video file
    
    Returns:
        List of paths to saved frames
    """
    paths = []
    base_name = Path(video_name).stem
    
    for i, frame in enumerate(frames):
        frame_path = os.path.join(TEMP_DIR, f"{base_name}_frame_{i:04d}.jpg")
        frame.save(frame_path)
        paths.append(frame_path)
    
    return paths

def analyze_frames_with_claude(frames, prompt):
    """
    Analyze a batch of frames with Claude
    
    Args:
        frames: List of PIL Image objects
        prompt: The text prompt for Claude
    
    Returns:
        Claude's response
    """
    # Prepare the multimodal message with frames
    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": prompt
                }
            ]
        }
    ]
    
    # Add frames to the message (up to 5 frames per message)
    for frame in frames:
        encoded_image = encode_image_for_claude(frame)
        messages[0]["content"].append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/jpeg",
                "data": encoded_image
            }
        })
    
    # Send request to Claude
    response = claude.messages.create(
        model="claude-3-opus-20240229",
        max_tokens=4000,
        messages=messages
    )
    
    return response.content[0].text

def create_vector_store_from_frames(frame_analysis_results, video_names, frame_timestamps):
    """
    Create a vector store from frame analysis results
    
    Args:
        frame_analysis_results: List of Claude's analysis of frames
        video_names: Names of the source videos
        frame_timestamps: Timestamps for frames in seconds
    
    Returns:
        FAISS vector store
    """
    documents = []
    
    # Create documents from frame analysis results
    for i, analysis in enumerate(frame_analysis_results):
        # Determine which video and timestamp this came from
        video_index = i // len(frame_timestamps[0])
        frame_index = i % len(frame_timestamps[0])
        
        if video_index < len(video_names) and frame_index < len(frame_timestamps[video_index]):
            video_name = video_names[video_index]
            timestamp = frame_timestamps[video_index][frame_index]
            
            # Format timestamp as minutes:seconds
            timestamp_formatted = f"{int(timestamp // 60)}:{int(timestamp % 60):02d}"
            
            # Create document with metadata
            doc = Document(
                page_content=analysis,
                metadata={
                    "source": video_name,
                    "timestamp": timestamp,
                    "timestamp_formatted": timestamp_formatted
                }
            )
            
            documents.append(doc)
    
    # Create vector store
    vector_store = FAISS.from_documents(documents, embeddings_model)
    return vector_store

def generate_final_response(query, vector_store, video_frames, all_video_names, all_frame_timestamps):
    """
    Generate a final response to the query using RAG with Claude
    
    Args:
        query: The user query
        vector_store: FAISS vector store of frame analyses
        video_frames: List of video frames as PIL Images
        all_video_names: List of video file names
        all_frame_timestamps: List of timestamps for frames
    
    Returns:
        Claude's response, key frames used, and relevant frames
    """
    # Retrieve relevant documents
    docs = vector_store.similarity_search(query, k=5)
    
    # Prepare content for Claude
    context = "\n\n".join([
        f"Source: {doc.metadata['source']} at {doc.metadata['timestamp_formatted']}\n{doc.page_content}" 
        for doc in docs
    ])
    
    # Track relevant frames from search results
    relevant_frames = []
    relevant_frame_info = []
    
    for doc in docs:
        video_name = doc.metadata['source']
        timestamp = doc.metadata['timestamp']
        
        # Find which video this is from
        video_index = -1
        for i, name in enumerate(all_video_names):
            if name == video_name:
                video_index = i
                break
        
        if video_index >= 0:
            # Find the closest frame to this timestamp
            timestamps = all_frame_timestamps[video_index]
            closest_frame_idx = min(range(len(timestamps)), 
                              key=lambda i: abs(timestamps[i] - timestamp))
            
            # Calculate the absolute index in the combined frames list
            absolute_idx = sum(len(ts) for ts in all_frame_timestamps[:video_index]) + closest_frame_idx
            
            if 0 <= absolute_idx < len(video_frames):
                relevant_frames.append(video_frames[absolute_idx])
                relevant_frame_info.append({
                    'video': video_name,
                    'timestamp': timestamp,
                    'timestamp_formatted': doc.metadata['timestamp_formatted']
                })
    
    # Select a few key frames to include with the prompt
    key_frame_indices = [0, min(5, len(video_frames)-1), min(10, len(video_frames)-1)]
    key_frames = [video_frames[i] for i in key_frame_indices if i < len(video_frames)]
    
    # Create prompt for Claude
    prompt = f"""
You are a knowledgeable assistant specializing in Symbox platform development workflows.

I need a comprehensive explanation about: {query}

Based on analysis of video tutorials, here's relevant information:

{context}

Please provide a detailed explanation of the complete journey for committing and deploying a Symbox component. 
Include all steps in the process, from initial development to final deployment.
Structure your answer with clear headings and a step-by-step guide that a developer can follow.
"""

    # Analyze with Claude
    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": prompt
                }
            ]
        }
    ]
    
    # Add key frames to the message
    for frame in key_frames:
        encoded_image = encode_image_for_claude(frame)
        messages[0]["content"].append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/jpeg",
                "data": encoded_image
            }
        })
    
    # Send request to Claude
    response = claude.messages.create(
        model="claude-3-opus-20240229",
        max_tokens=4000,
        messages=messages
    )
    
    return response.content[0].text, key_frames, relevant_frames, relevant_frame_info

def save_response_as_docx(response_text, key_frames, relevant_frames, relevant_frame_info, query):
    """
    Save the response text and frames as a docx file
    
    Args:
        response_text: Text response from Claude
        key_frames: List of key frames as PIL Images
        relevant_frames: List of relevant frames from search results
        relevant_frame_info: Metadata about relevant frames
        query: The original query
    
    Returns:
        Path to the saved docx file
    """
    # Create a new document
    doc = docx.Document()
    
    # Set document properties
    doc.core_properties.title = f"Symbox Guide: {query}"
    doc.core_properties.author = "Symbox Documentation Generator"
    doc.core_properties.created = datetime.now()
    
    # Add a cover page
    doc.add_heading(f"Symbox Documentation", level=0)
    doc.add_heading(query, level=1)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Add a page break after the cover
    doc.add_page_break()
    
    # Add table of contents heading
    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    run = toc_para.add_run("(Right-click and select 'Update Field' to update the table of contents)")
    run.italic = True
    run.font.size = docx.shared.Pt(9)
    
    # Add TOC field
    para = doc.add_paragraph()
    run = para.add_run()
    fld_char = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar', 
                                        {'fldCharType': 'begin'})
    run._element.append(fld_char)
    
    instrText = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText', {})
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)
    
    fld_char = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar', 
                                        {'fldCharType': 'separate'})
    run._element.append(fld_char)
    
    # Add static TOC placeholder text
    run = para.add_run("Table of Contents placeholder. Right-click and select 'Update Field' to update.")
    
    fld_char = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar', 
                                        {'fldCharType': 'end'})
    run._element.append(fld_char)
    
    # Add a page break after TOC
    doc.add_page_break()
    
    # Add content heading
    doc.add_heading("Guide Content", level=1)
    
    # Process the response text to handle markdown formatting
    paragraphs = response_text.split('\n\n')
    i = 0
    
    while i < len(paragraphs):
        para = paragraphs[i].strip()
        
        # Check if paragraph is a heading (starts with #)
        if para.startswith('#'):
            # Count the number of # to determine heading level
            heading_level = 0
            for char in para:
                if char == '#':
                    heading_level += 1
                else:
                    break
            
            heading_level = min(heading_level, 9)  # Word only supports heading levels 1-9
            heading_text = para[heading_level:].strip()
            doc.add_heading(heading_text, level=heading_level)
        
        # Check if paragraph is a bulleted list (starts with * or -)
        elif para.startswith('* ') or para.startswith('- '):
            # Handle multi-line bulleted list
            bullet_items = [para]
            j = i + 1
            while j < len(paragraphs) and (paragraphs[j].strip().startswith('* ') or 
                                          paragraphs[j].strip().startswith('- ')):
                bullet_items.append(paragraphs[j].strip())
                j += 1
            
            for item in bullet_items:
                # Remove the bullet marker and add as a list item
                item_text = item[2:].strip()
                doc.add_paragraph(item_text, style='List Bullet')
            
            i = j - 1  # Update the counter to skip processed items
        
        # Check if paragraph is a numbered list (starts with 1., 2., etc.)
        elif para[0].isdigit() and len(para) > 1 and para[1:].startswith('. '):
            # Handle multi-line numbered list
            number_items = [para]
            current_number = int(para[0])
            j = i + 1
            
            while j < len(paragraphs):
                next_para = paragraphs[j].strip()
                if next_para.startswith(f"{current_number + 1}. "):
                    number_items.append(next_para)
                    current_number += 1
                    j += 1
                else:
                    break
            
            for item in number_items:
                # Find the position after the number and period
                pos = item.find('. ')
                if pos != -1:
                    item_text = item[pos + 2:].strip()
                    doc.add_paragraph(item_text, style='List Number')
            
            i = j - 1  # Update the counter to skip processed items
            
        # Regular paragraph
        else:
            if para:  # Only add non-empty paragraphs
                doc.add_paragraph(para)
        
        i += 1
    
    # Add a page break before images
    doc.add_page_break()
    
    # Add a section for relevant frames from the search
    if relevant_frames:
        doc.add_heading("Relevant Tutorial Screenshots", level=1)
        
        for i, (frame, info) in enumerate(zip(relevant_frames, relevant_frame_info)):
            # Save the image to a temporary file
            temp_img_path = os.path.join(TEMP_DIR, f"relevant_frame_{i}.jpg")
            frame.save(temp_img_path)
            
            # Add image to document with caption
            caption = f"From: {os.path.basename(info['video'])} at {info['timestamp_formatted']}"
            doc.add_paragraph(caption)
            doc.add_picture(temp_img_path, width=Inches(6))
    
    # Add a section for key frames used in the prompt
    doc.add_heading("Additional Tutorial Screenshots", level=1)
    
    # Save key frames to temporary files and add to document
    for i, frame in enumerate(key_frames):
        # Save the image to a temporary file
        temp_img_path = os.path.join(TEMP_DIR, f"key_frame_{i}.jpg")
        frame.save(temp_img_path)
        
        # Add image to document with caption
        doc.add_paragraph(f"Frame {i+1}")
        doc.add_picture(temp_img_path, width=Inches(6))
    
    # Save the document
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    docx_path = f"response_{timestamp}.docx"
    doc.save(docx_path)
    
    return docx_path

def main():
    """Main function to orchestrate the multimodal RAG workflow"""
    print(f"Starting multimodal RAG analysis on {len(VIDEO_FILES)} video files...")
    
    all_frames = []
    all_video_names = []
    all_frame_timestamps = []
    
    # Extract frames from all videos
    for video_file in VIDEO_FILES:
        print(f"Processing video: {video_file}")
        frames, timestamps = extract_frames_from_video(video_file)
        all_frames.extend(frames)
        all_video_names.append(video_file)
        all_frame_timestamps.append(timestamps)
    
    # Analyze frames in batches (Claude can handle up to 5 images per request)
    frame_analysis_results = []
    batch_size = 5
    
    for i in range(0, len(all_frames), batch_size):
        print(f"Analyzing batch {i // batch_size + 1}/{(len(all_frames) + batch_size - 1) // batch_size}")
        batch_frames = all_frames[i:i+batch_size]
        
        analysis_prompt = """
        Analyze these video frames from a tutorial about Symbox component development and deployment.
        Describe what you see in detail, focusing on:
        1. Any visible user interface elements and what they represent
        2. Any text or code visible in the frame
        3. Any workflow steps being demonstrated
        4. Any version control, commit, or deployment processes shown
        
        Be specific and detailed in your observations.
        """
        
        analysis = analyze_frames_with_claude(batch_frames, analysis_prompt)
        frame_analysis_results.append(analysis)
        
        # Save to prevent losing progress
        with open(f"frame_analysis_batch_{i // batch_size}.txt", "w") as f:
            f.write(analysis)
        
        # Respect API rate limits
        time.sleep(1)
    
    # Create vector store from frame analyses
    print("Creating vector store from frame analyses...")
    vector_store = create_vector_store_from_frames(
        frame_analysis_results, 
        all_video_names, 
        all_frame_timestamps
    )
    
    # Generate final response
    print(f"Generating final response to query: {QUERY}")
    response, key_frames, relevant_frames, relevant_frame_info = generate_final_response(
        QUERY, vector_store, all_frames, all_video_names, all_frame_timestamps
    )
    
    # Save the result as text file
    output_file = f"response_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    with open(output_file, "w") as f:
        f.write(response)
    
    # Save the result as docx file with images
    docx_path = save_response_as_docx(response, key_frames, relevant_frames, relevant_frame_info, QUERY)
    
    print(f"Response saved to {output_file}")
    print(f"Response with images saved to {docx_path}")
    print("\nFinal Response:")
    print("-" * 80)
    print(response)
    print("-" * 80)

if __name__ == "__main__":
    main()
