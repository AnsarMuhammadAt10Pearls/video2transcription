import os
import docx
import glob
import re
from collections import Counter
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.schema import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Simple word tokenization without NLTK
def simple_tokenize(text):
    """Simple word tokenization without external dependencies"""
    # Convert to lowercase and split by non-alphanumeric characters
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
    # Filter out common stopwords
    stopwords = {'the', 'and', 'are', 'for', 'was', 'not', 'but', 'this', 'that', 'with', 
                'from', 'have', 'has', 'had', 'they', 'will', 'what', 'when', 'where', 'who',
                'which', 'how', 'why', 'can', 'could', 'would', 'should', 'than', 'then', 'them',
                'their', 'there', 'these', 'those', 'some', 'such', 'into', 'about', 'been',
                'you', 'your', 'all', 'any', 'own'}
    return [word for word in words if word not in stopwords]

def extract_text_from_docx(file_path):
    """Extract text content from a Word document"""
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def get_related_images(document_name):
    """Find images that might be related to a document based on file name"""
    document_base = os.path.splitext(os.path.basename(document_name))[0]
    image_path = "images"
    
    # Make sure the images directory exists
    if not os.path.exists(image_path):
        os.makedirs(image_path)
        return []
    
    # Look for images with the document name in their filename
    related_images = []
    for ext in ['*.jpg', '*.jpeg', '*.png', '*.gif']:
        for img_file in glob.glob(os.path.join(image_path, ext)):
            img_basename = os.path.basename(img_file)
            # Check if the document name is part of the image name
            if document_base.lower() in img_basename.lower():
                related_images.append(img_file)
    
    # Add all images if none are directly linked
    if not related_images:
        for ext in ['*.jpg', '*.jpeg', '*.png', '*.gif']:
            related_images.extend(glob.glob(os.path.join(image_path, ext)))
    
    return related_images

def process_documents(file_paths):
    """Process multiple document files into a vector database"""
    documents = []
    
    # Extract text from each document
    for file_path in file_paths:
        text = extract_text_from_docx(file_path)
        
        # Find related images
        related_images = get_related_images(file_path)
        
        # Create a Document object with metadata including the source file and related images
        doc = Document(
            page_content=text, 
            metadata={
                "source": file_path,
                "related_images": related_images
            }
        )
        documents.append(doc)
    
    # Split documents into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    chunks = text_splitter.split_documents(documents)
    
    # Create vector embeddings and store in a vector database
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vector_store = FAISS.from_documents(chunks, embeddings)
    
    return vector_store

def extract_keywords(text):
    """Extract meaningful keywords from text without external NLP libraries"""
    # Use simple tokenization instead of NLTK
    tokens = simple_tokenize(text)
    
    # Count frequencies
    word_freq = Counter(tokens)
    
    # Get keywords (words occurring at least once)
    keywords = {word for word, count in word_freq.items() if len(word) > 2}
    
    return keywords

def calculate_content_similarity(query_keywords, doc_content):
    """Calculate semantic similarity between query and document content"""
    # Extract words from document content
    doc_words = simple_tokenize(doc_content)
    
    # Convert to sets for intersection
    doc_keywords = set(doc_words)
    
    # Calculate keyword overlap
    if not query_keywords or not doc_keywords:
        return 0.0
        
    # Calculate Jaccard similarity (intersection over union)
    intersection = query_keywords.intersection(doc_keywords)
    union = query_keywords.union(doc_keywords)
    
    if not union:
        return 0.0
    
    return len(intersection) / len(union)

def find_top_matching_images(query, source_docs, max_images=3):
    """Find the top images that best match the query and context"""
    # Extract meaningful keywords from query
    query_keywords = extract_keywords(query)
    
    # Collect all images and calculate relevance scores
    image_scores = {}
    
    # First, check if there are any images in the source documents
    all_images = set()
    for doc in source_docs:
        if "related_images" in doc.metadata and doc.metadata["related_images"]:
            all_images.update(doc.metadata["related_images"])
    
    if not all_images:
        return []
    
    # Now score each image
    for doc_idx, doc in enumerate(source_docs):
        if "related_images" not in doc.metadata or not doc.metadata["related_images"]:
            continue
        
        # Calculate content similarity between query and document
        content_similarity = calculate_content_similarity(query_keywords, doc.page_content)
        
        # Apply position weight (earlier documents are more relevant)
        position_weight = 1.0 / (doc_idx + 1)
        
        for img in doc.metadata["related_images"]:
            # Extract words from image filename
            img_basename = os.path.basename(img)
            img_name = os.path.splitext(img_basename)[0].replace('_', ' ').replace('-', ' ')
            img_keywords = extract_keywords(img_name)
            
            # Calculate image name similarity to query
            img_name_similarity = 0
            if query_keywords and img_keywords:
                overlap = query_keywords.intersection(img_keywords)
                img_name_similarity = len(overlap) / len(query_keywords) if query_keywords else 0
            
            # Calculate overall score with weighted components
            # Higher weight for content similarity, lower for filename and position
            score = (content_similarity * 5.0) + (img_name_similarity * 2.0) + (position_weight * 1.0)
            
            # Add random small variation to avoid ties (0-0.1)
            score += (hash(img) % 100) / 1000.0
            
            # Store the image with its score and source
            if img not in image_scores or score > image_scores[img][0]:
                image_scores[img] = (score, doc.metadata["source"])
    
    # Sort images by score (descending)
    sorted_images = sorted(
        [(img, score, source) for img, (score, source) in image_scores.items()],
        key=lambda x: x[1],
        reverse=True
    )
    
    # Return top N images
    return sorted_images[:max_images]

def create_qa_chain(vector_store):
    """Create a question-answering chain using the vector store"""
    # Initialize the Groq language model
    llm = ChatGroq(temperature=0.7, model_name="llama3-8b-8192")
    
    # Create the retriever as a runnable
    retriever = vector_store.as_retriever(search_kwargs={"k": 3})
    
    # Define a proper ChatPromptTemplate for the LLM
    prompt = ChatPromptTemplate.from_template("""
    Answer the following question based on the provided context:
    
    Context:
    {context}
    
    Question: {question}
    
    Answer:
    """)
    
    # Create a modern LCEL chain with proper formatting
    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)
    
    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    return rag_chain, retriever

def main():
    # Paths to your Word documents
    file_paths = ["first 5 min.docx", "next 5 min.docx"]
    
    # Check if files exist
    for file_path in file_paths:
        if not os.path.exists(file_path):
            print(f"Error: File '{file_path}' not found.")
            return
    
    print("Processing documents...")
    vector_store = process_documents(file_paths)
    qa_chain, retriever = create_qa_chain(vector_store)
    
    print("\nDocuments processed successfully. You can now ask questions about their content.")
    print("Type 'exit' to quit.")
    
    # Default question
    default_question = "Complete Journey of Committing and Deploying of a Symbox Component"
    
    while True:
        # Show default question in prompt
        user_input = input(f"\nQuestion (press Enter for default or type 'exit' to quit): [{default_question}] ")
        
        # Use default question if user just pressed Enter
        if user_input.strip() == "":
            query = default_question
            print(f"Using default question: {default_question}")
        else:
            query = user_input
            
        if query.lower() == 'exit':
            break
        
        # Use invoke() method instead of get_relevant_documents
        source_docs = retriever.invoke(query)
        
        # Get the answer using the updated LCEL chain
        answer = qa_chain.invoke(query)
        
        # Display the answer
        print("\nAnswer:", answer)
        
        # Display sources
        print("\nSources:")
        for doc in source_docs:
            print(f"- {doc.metadata['source']}")
            
        # Find and display top 3 most relevant images (without debug output)
        top_images = find_top_matching_images(query, source_docs, max_images=3)
        
        if top_images:
            print("\nTop Related Images:")
            for img, score, source in top_images:
                print(f"- {img} (from {source})")
        else:
            print("\nNo related images found.")

if __name__ == "__main__":
    main()