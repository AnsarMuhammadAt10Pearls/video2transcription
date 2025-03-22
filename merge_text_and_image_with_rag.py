import os
import glob
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
import numpy as np
from PIL import Image
from PIL.ExifTags import TAGS
import time
from tabulate import tabulate
import base64
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough

# Load environment variables
load_dotenv()

# Hard-coded query
QUERY = "Complete Journey of Committing and Deploying of a Symbox Component"

# Initialize embeddings model
embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Define Groq models to use
GROQ_MODEL = "llama3-70b-8192"
GROQ_TEMP = 0.1

# Check if OCR is available
OCR_AVAILABLE = False
try:
    import pytesseract
    import cv2
    # Try to set tesseract path for Windows users
    pytesseract.pytesseract.tesseract_cmd = r'd:\tesseract\tesseract.exe'
    # Test if tesseract is installed
    pytesseract.get_tesseract_version()
    OCR_AVAILABLE = True
    print("OCR is available and will be used to extract text from images")
except Exception as e:
    print(f"OCR is not available: {str(e)}")
    print("Image search will work but without OCR text extraction.")

# Create a cache directory for OCR results
OCR_CACHE_DIR = "ocr_cache"
os.makedirs(OCR_CACHE_DIR, exist_ok=True)

def extract_text_from_docx(file_path):
    """Extract text content from a Word document"""
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def extract_text_from_image(image_path):
    """Extract text from image using OCR with caching to avoid repeated processing"""
    if not OCR_AVAILABLE:
        return ""
        
    # Generate a simple hash of the image path for caching
    import hashlib
    img_hash = hashlib.md5(image_path.encode()).hexdigest()
    
    cache_file = os.path.join(OCR_CACHE_DIR, f"{img_hash}.txt")
    
    # Check if we have this image's OCR results in cache
    if os.path.exists(cache_file):
        with open(cache_file, 'r', encoding='utf-8') as f:
            return f.read()
    
    # Process the image with OCR
    try:
        # Use CV2 to preprocess the image for better OCR results
        img = cv2.imread(image_path)
        if img is None:
            return ""
            
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Apply threshold to get black and white image
        _, threshold = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Perform OCR on the threshold version
        text = pytesseract.image_to_string(threshold)
        
        # Cache the results
        with open(cache_file, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return text
    except Exception as e:
        print(f"OCR error for {image_path}: {str(e)}")
        return ""

def extract_image_metadata(image_path):
    """Extract meaningful metadata from an image when possible"""
    try:
        with Image.open(image_path) as img:
            # Get EXIF data
            exif_data = {}
            if hasattr(img, '_getexif') and img._getexif():
                for tag, value in img._getexif().items():
                    if tag in TAGS:
                        exif_data[TAGS[tag]] = value
                        
            # Extract description or title if available
            description = exif_data.get('ImageDescription', '')
            
            # Basic image properties
            properties = {
                'format': img.format,
                'size': img.size,
                'mode': img.mode,
                'description': description,
                'filename': os.path.basename(image_path)
            }
            
            return properties
    except Exception as e:
        # If any error occurs, return minimal information
        return {'format': 'unknown', 'error': str(e), 'filename': os.path.basename(image_path)}

def get_related_images(document_name):
    """Find images that might be related to a document based on file name"""
    document_base = os.path.splitext(os.path.basename(document_name))[0]
    image_path = "images"
    
    # Make sure the images directory exists
    if not os.path.exists(image_path):
        os.makedirs(image_path)
        return []
    
    # Look for images with the document name in their filename
    related_images = []
    for ext in ['*.jpg', '*.jpeg', '*.png', '*.gif']:
        for img_file in glob.glob(os.path.join(image_path, ext)):
            img_basename = os.path.basename(img_file)
            # Check if the document name is part of the image name
            if document_base.lower() in img_basename.lower():
                related_images.append(img_file)
    
    # Add all images if none are directly linked
    if not related_images:
        for ext in ['*.jpg', '*.jpeg', '*.png', '*.gif']:
            related_images.extend(glob.glob(os.path.join(image_path, ext)))
    
    return related_images

def process_documents(file_paths):
    """Process multiple document files into a vector database"""
    documents = []
    
    # Extract text from each document
    for file_path in file_paths:
        text = extract_text_from_docx(file_path)
        
        # Find related images
        related_images = get_related_images(file_path)
        
        # Create a Document object with metadata including the source file and related images
        doc = Document(
            page_content=text, 
            metadata={
                "source": file_path,
                "related_images": related_images
            }
        )
        documents.append(doc)
    
    # Split documents into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    chunks = text_splitter.split_documents(documents)
    
    # Create vector embeddings and store in a vector database
    vector_store = FAISS.from_documents(chunks, embeddings_model)
    
    return vector_store

def create_image_embedding(image_path, doc_content, query):
    """Create a semantic embedding for an image based on its filename, OCR text, and document context"""
    # Extract image filename and path components
    basename = os.path.basename(image_path)
    name_without_ext = os.path.splitext(basename)[0]
    
    # Clean up the name by replacing separators with spaces
    cleaned_name = name_without_ext.replace('_', ' ').replace('-', ' ')
    
    # Extract document context (first 300 chars for relevance)
    doc_preview = doc_content[:300] if len(doc_content) > 300 else doc_content
    
    # Try to get image metadata
    try:
        metadata = extract_image_metadata(image_path)
        img_desc = metadata.get('description', '')
    except:
        img_desc = ""
    
    # Get OCR text if available
    ocr_text = ""
    has_exact_match = False
    if OCR_AVAILABLE:
        ocr_text = extract_text_from_image(image_path)
        # Check if query appears in OCR text
        if query.lower() in ocr_text.lower():
            has_exact_match = True
    
    # Combine all relevant text for embedding
    combined_text = f"Image: {cleaned_name}. Context: {doc_preview}. {img_desc}. "
    if ocr_text:
        combined_text += f"OCR Text: {ocr_text}"
    
    # Generate embedding
    return embeddings_model.embed_query(combined_text), ocr_text, has_exact_match

def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors"""
    dot_product = np.dot(vec1, vec2)
    norm_a = np.linalg.norm(vec1)
    norm_b = np.linalg.norm(vec2)
    if norm_a == 0 or norm_b == 0:
        return 0
    return dot_product / (norm_a * norm_b)

def find_top_matching_images(query, source_docs, max_images=3, min_similarity_threshold=0.5):
    """Find the top images using AI embeddings for semantic matching with a minimum threshold"""
    # Skip if no documents retrieved
    if not source_docs:
        return []
        
    # Generate embedding for the query
    query_embedding = embeddings_model.embed_query(query)
    
    # Collect all images and their metadata
    images_with_embeddings = []
    
    # Map to prevent duplicate processing
    processed_images = {}
    
    # First check if we have any images
    all_images = set()
    for doc in source_docs:
        if "related_images" in doc.metadata and doc.metadata["related_images"]:
            all_images.update(doc.metadata["related_images"])
    
    if not all_images:
        return []
    
    print(f"Found {len(all_images)} potentially related images, analyzing...")
    
    # Create embeddings for each image based on its document context
    for doc_idx, doc in enumerate(source_docs):
        if "related_images" not in doc.metadata or not doc.metadata["related_images"]:
            continue
            
        for img_path in doc.metadata["related_images"]:
            # Skip if already processed
            if img_path in processed_images:
                continue
                
            # Create embedding that combines image filename with document context and OCR
            img_embedding, ocr_text, has_exact_match = create_image_embedding(img_path, doc.page_content, query)
            
            # Store for similarity calculation
            images_with_embeddings.append({
                'path': img_path,
                'embedding': img_embedding,
                'source': doc.metadata["source"],
                'doc_idx': doc_idx,
                'ocr_text': ocr_text,
                'has_exact_match': has_exact_match
            })
            
            # Mark as processed
            processed_images[img_path] = True
    
    # Calculate similarity scores
    scored_images = []
    for img_data in images_with_embeddings:
        # Calculate semantic similarity between query and image context
        similarity = cosine_similarity(query_embedding, img_data['embedding'])
        
        # Apply boost for exact matches found in OCR
        if img_data.get('has_exact_match', False):
            similarity = min(1.0, similarity * 1.5)  # Boost by 50% but cap at 1.0
            print(f"Found exact match in {os.path.basename(img_data['path'])}! Boosting similarity.")
        
        # Add position boost (earlier docs get higher priority)
        position_boost = 0.05 * (len(source_docs) - img_data['doc_idx']) / len(source_docs)
        
        # Final score is similarity plus small position boost
        final_score = similarity + position_boost
        
        # Only include images that meet the minimum similarity threshold
        if similarity >= min_similarity_threshold:
            scored_images.append({
                'path': img_data['path'],
                'score': final_score,
                'source': img_data['source'],
                'ocr_text': img_data.get('ocr_text', ''),
                'has_exact_match': img_data.get('has_exact_match', False)
            })
    
    # Sort by score (descending)
    sorted_images = sorted(scored_images, key=lambda x: x['score'], reverse=True)
    
    # Return top N (or fewer if not enough meet the threshold)
    return sorted_images[:max_images]

def encode_image(image_path):
    """Encode image to base64 string for HTML display"""
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

def create_qa_chain(vector_store):
    """Create a question-answering chain using the vector store"""
    # Initialize the Groq language model
    llm = ChatGroq(temperature=GROQ_TEMP, model_name=GROQ_MODEL)
    
    # Create the retriever as a runnable
    retriever = vector_store.as_retriever(search_kwargs={"k": 4})
    
    return llm, retriever

def generate_answer_with_images(query, llm, source_docs, top_images):
    """Generate final answer with text and images"""
    # Format documents for context
    formatted_docs = ""
    for i, doc in enumerate(source_docs):
        formatted_docs += f"Document {i+1} (from {doc.metadata['source']}):\n{doc.page_content}\n\n"
    
    # Format images for context
    image_context = ""
    for i, img in enumerate(top_images):
        img_name = os.path.basename(img['path'])
        img_source = os.path.basename(img['source'])
        ocr_preview = ""
        if img.get('ocr_text', ''):
            ocr_preview = f" Text content: {img['ocr_text'][:200]}..."
        image_context += f"Image {i+1}: {img_name} (from {img_source}).{ocr_preview}\n"
    
    # Create prompt for comprehensive answer
    prompt = f"""
    You are an expert in software development processes and deployment systems.
    
    Please provide a comprehensive answer to this question:
    "{query}"
    
    Use the following document context:
    {formatted_docs}
    
    The following images also appear to be relevant:
    {image_context}
    
    Instructions:
    1. Provide a well-structured, detailed answer covering all stages of the process
    2. Include specific steps and procedures mentioned in the documents
    3. Reference specific images when they help illustrate a step or concept
    4. Format your answer with clear sections and bullet points where appropriate
    5. If any steps are unclear or information is missing, note this briefly
    
    Your response should be comprehensive yet easy to follow, focusing on the complete process of committing and deploying a Symbox component.
    """
    
    # Get response from LLM
    response = llm.invoke(prompt)
    
    # Format as a final comprehensive answer with image references
    return response.content

def create_html_report(query, answer, source_docs, top_images):
    """Create an HTML report with the answer and images"""
    
    # Prepare HTML for images
    images_html = ""
    for i, img in enumerate(top_images):
        img_path = img['path']
        img_name = os.path.basename(img_path)
        img_base64 = encode_image(img_path)
        img_score = img['score']
        img_source = os.path.basename(img['source'])
        
        has_match = "Yes" if img.get('has_exact_match', False) else "No"
        
        images_html += f"""
        <div class="image-card">
            <h3>Image {i+1}: {img_name}</h3>
            <p>Source: {img_source} | Relevance Score: {img_score:.4f} | Contains Exact Match: {has_match}</p>
            <img src="data:image/jpeg;base64,{img_base64}" alt="{img_name}" class="result-image">
        </div>
        """
    
    # Prepare HTML for document sources
    sources_html = ""
    for i, doc in enumerate(source_docs):
        source_name = doc.metadata['source']
        sources_html += f"<li>{source_name}</li>"
    
    # Convert answer text to HTML (replace newlines with <br>)
    answer_html = answer.replace('\n', '<br>')
    
    # Create the complete HTML report
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Symbox Deployment Process</title>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }}
            .header {{ background-color: #f2f2f2; padding: 20px; border-radius: 5px; margin-bottom: 20px; }}
            .answer-section {{ margin-bottom: 30px; }}
            .images-section {{ display: flex; flex-wrap: wrap; gap: 20px; margin-bottom: 30px; }}
            .image-card {{ border: 1px solid #ddd; border-radius: 5px; padding: 15px; width: 100%; max-width: 600px; }}
            .result-image {{ max-width: 100%; height: auto; margin-top: 10px; }}
            .sources-section {{ background-color: #f9f9f9; padding: 15px; border-radius: 5px; }}
            h1, h2, h3 {{ color: #333; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>Symbox Deployment Process Analysis</h1>
            <p><strong>Query:</strong> {query}</p>
        </div>
        
        <div class="answer-section">
            <h2>Complete Answer:</h2>
            <div>{answer_html}</div>
        </div>
        
        <div class="images-section">
            <h2>Relevant Images:</h2>
            {images_html}
        </div>
        
        <div class="sources-section">
            <h2>Document Sources:</h2>
            <ul>
                {sources_html}
            </ul>
        </div>
    </body>
    </html>
    """
    
    # Write HTML to file
    with open("symbox_deployment_process.html", "w", encoding="utf-8") as f:
        f.write(html)
    
    return "symbox_deployment_process.html"

def main():
    print(f"Starting comprehensive analysis for query: '{QUERY}'")
    start_time = time.time()
    
    # 1. Process documents
    file_paths = ["first 5 min.docx", "next 5 min.docx"]
    
    # Check if files exist
    for file_path in file_paths:
        if not os.path.exists(file_path):
            print(f"Error: File '{file_path}' not found.")
            return
    
    print("Processing documents...")
    vector_store = process_documents(file_paths)
    llm, retriever = create_qa_chain(vector_store)
    
    # 2. Retrieve relevant document chunks
    print("Retrieving relevant information...")
    source_docs = retriever.invoke(QUERY)
    
    # 3. Find top matching images
    print("Finding relevant images...")
    top_images = find_top_matching_images(QUERY, source_docs, max_images=6, min_similarity_threshold=0.3)
    
    # 4. Generate comprehensive answer
    print("Generating comprehensive answer...")
    answer = generate_answer_with_images(QUERY, llm, source_docs, top_images)
    
    # 5. Create HTML report with embedded images
    print("Creating HTML report...")
    html_path = create_html_report(QUERY, answer, source_docs, top_images)
    
    end_time = time.time()
    execution_time = end_time - start_time
    
    print(f"\nProcess completed in {execution_time:.2f} seconds.")
    print(f"HTML report saved to: {html_path}")
    print("\nAnswer summary:")
    print("="*80)
    # Print first 400 characters of answer as preview
    print(answer[:400] + "..." if len(answer) > 400 else answer)
    print("="*80)
    print(f"View full report with images in {html_path}")

if __name__ == "__main__":
    main()
